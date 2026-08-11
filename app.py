import streamlit as st
import time
import io
from docx import Document
from agents import build_reader_agent, build_search_agent, writer_chain, critic_chain

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResearchFlow · AI Research Agent",
    page_icon="🔬",
    layout="wide",
)

# ── Light CSS (just a bit of polish, not a full theme) ────────────────────────
st.markdown("""
<style>
.stButton > button {
    border-radius: 8px;
    font-weight: 600;
}
.section-heading {
    font-size: 1.3rem;
    font-weight: 700;
    margin: 1.5rem 0 0.75rem;
}
.notice {
    font-size: 0.75rem;
    color: #888;
    text-align: center;
    margin-top: 3rem;
}
</style>
""", unsafe_allow_html=True)


# ── Helper: build a Word document from the report + feedback ─────────────────
def build_docx(topic: str, report: str, feedback: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Research Report: {topic}", level=1)

    doc.add_heading("Report", level=2)
    for line in report.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.add_heading("Critic Feedback", level=2)
    for line in feedback.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Session state init ────────────────────────────────────────────────────────
for key in ("results", "running", "done"):
    if key not in st.session_state:
        st.session_state[key] = {} if key == "results" else False


# ── Header ────────────────────────────────────────────────────────────────────
st.title("🔬 ResearFlow")
st.caption("Four AI agents collaborate — searching, scraping, writing, and critiquing — to deliver a research report on any topic.")

st.divider()


# ── Input ─────────────────────────────────────────────────────────────────────
topic = st.text_input(
    "Research Topic",
    placeholder="e.g. Quantum computing breakthroughs in 2025",
    key="topic_input",
)
run_btn = st.button(" Run Research Pipeline", use_container_width=True)

st.caption("Try: LLM agents 2025 · CRISPR gene editing · Fusion energy progress")


# ── Pipeline status ───────────────────────────────────────────────────────────
st.markdown('<div class="section-heading">Pipeline</div>', unsafe_allow_html=True)

r = st.session_state.results
steps = [
    ("search", " Search Agent", "Gathers recent web information"),
    ("reader", " Reader Agent", "Scrapes & extracts deep content"),
    ("writer", " Writer Chain", "Drafts the full research report"),
    ("critic", " Critic Chain", "Reviews & scores the report"),
]

cols = st.columns(4)
for i, (key, title, desc) in enumerate(steps):
    with cols[i]:
        if key in r:
            st.success(f"{title}\n\n✓ Done")
        elif st.session_state.running:
            st.info(f"{title}\n\n● Running…")
        else:
            st.caption(f"{title}\n\nWaiting")
        st.caption(desc)


# ── Run pipeline ──────────────────────────────────────────────────────────────
if run_btn:
    if not topic.strip():
        st.warning("Please enter a research topic first.")
    else:
        st.session_state.results = {}
        st.session_state.running = True
        st.session_state.done = False
        st.rerun()

if st.session_state.running and not st.session_state.done:
    results = {}
    topic_val = st.session_state.topic_input

    with st.spinner("🔍 Search Agent is working…"):
        search_agent = build_search_agent()
        sr = search_agent.invoke({
            "messages": [("user", f"Find recent, reliable and detailed information about: {topic_val}")]
        })
        results["search"] = sr["messages"][-1].content
        st.session_state.results = dict(results)

    with st.spinner("📄 Reader Agent is scraping top resources…"):
        reader_agent = build_reader_agent()
        rr = reader_agent.invoke({
            "messages": [("user",
                f"Based on the following search results about '{topic_val}', "
                f"pick the most relevant URL and scrape it for deeper content.\n\n"
                f"Search Results:\n{results['search'][:800]}"
            )]
        })
        results["reader"] = rr["messages"][-1].content
        st.session_state.results = dict(results)

    with st.spinner("✍️ Writer is drafting the report…"):
        research_combined = (
            f"SEARCH RESULTS:\n{results['search']}\n\n"
            f"DETAILED SCRAPED CONTENT:\n{results['reader']}"
        )
        results["writer"] = writer_chain.invoke({
            "topic": topic_val,
            "research": research_combined
        })
        st.session_state.results = dict(results)

    with st.spinner(" Critic is reviewing the report…"):
        results["critic"] = critic_chain.invoke({
            "report": results["writer"]
        })
        st.session_state.results = dict(results)

    st.session_state.running = False
    st.session_state.done = True
    st.rerun()


# ── Results display ───────────────────────────────────────────────────────────
r = st.session_state.results

if r:
    st.divider()
    st.markdown('<div class="section-heading">Results</div>', unsafe_allow_html=True)

    if "search" in r:
        with st.expander(" Search Results (raw)"):
            st.write(r["search"])

    if "reader" in r:
        with st.expander(" Scraped Content (raw)"):
            st.write(r["reader"])

    if "writer" in r:
        st.subheader("Final Research Report")
        st.markdown(r["writer"])

    if "critic" in r:
        st.subheader(" Critic Feedback")
        st.markdown(r["critic"])

    if "writer" in r and "critic" in r:
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="Download as Markdown (.md)",
                data=r["writer"],
                file_name=f"research_report_{int(time.time())}.md",
                mime="text/markdown",
                use_container_width=True,
            )
        with col2:
            docx_bytes = build_docx(st.session_state.topic_input, r["writer"], r["critic"])
            st.download_button(
                label="⬇ Download as Word (.docx)",
                data=docx_bytes,
                file_name=f"research_report_{int(time.time())}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown('<div class="notice">ResearchMind · Powered by LangChain multi-agent pipeline · Built with Streamlit</div>', unsafe_allow_html=True)